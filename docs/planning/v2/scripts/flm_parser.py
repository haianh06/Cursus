"""
FLM -> RAG ingestion parser
----------------------------
Đọc file .docx export từ FLM (dạng "FPT_University_Learning_Materials_Formatted.docx"
cho curriculum, hoặc "<MaMon>_Formatted_Tables_Preserved.docx" cho từng syllabus)
và xuất ra JSON đã chunk sẵn kèm metadata, đúng format cần cho pgvector/Qdrant ingest
(FR-2.1 trong SRS: mỗi chunk phải kèm tên môn, mã môn, mục cụ thể để trích nguồn).

Cách dùng:
    python flm_parser.py curriculum path/to/curriculum.docx -o courses.json
    python flm_parser.py syllabus path/to/SSA101.docx -o chunks_SSA101.json
"""
import sys, json, re
from docx import Document


def table_to_kv(table):
    """Bảng 2 cột Field/Value -> dict."""
    kv = {}
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2 and cells[0]:
            key = re.sub(r"[*:]", "", cells[0]).strip()
            if key and key not in kv:
                kv[key] = cells[1]
    return kv


def table_to_records(table):
    """Bảng có header hàng đầu -> list[dict]."""
    rows = [[c.text.strip() for c in r.cells] for r in table.rows]
    if not rows:
        return []
    header = [re.sub(r"\s+", " ", h).strip() for h in rows[0]]
    records = []
    for r in rows[1:]:
        if len(r) != len(header):
            continue
        rec = {header[i]: r[i] for i in range(len(header)) if header[i]}
        if any(v for v in rec.values()):
            records.append(rec)
    return records


def parse_curriculum(path):
    doc = Document(path)
    tables = doc.tables
    meta = {}
    subjects = []
    for t in tables:
        header_row = [c.text.strip() for c in t.rows[0].cells] if t.rows else []
        joined = " ".join(header_row)
        if "Field" in joined and "Value" in joined:
            meta.update(table_to_kv(t))
        elif "Subject Code" in joined and "Subject Name" in joined:
            subjects = table_to_records(t)
    return {"curriculum_meta": meta, "subjects": subjects, "subject_count": len(subjects)}


def parse_syllabus(path):
    doc = Document(path)
    tables = doc.tables
    meta = {}
    clos = []
    sessions = []
    for t in tables:
        header_row = [c.text.strip() for c in t.rows[0].cells] if t.rows else []
        joined = " ".join(header_row)
        if "Session" in joined and "Topic" in joined:
            sessions = table_to_records(t)
        elif "CLO" in joined and ("Details" in joined or "CLO Details" in joined):
            clos = table_to_records(t)
        elif re.search(r"Syllabus (ID|Name)", joined) or "Subject Code" in joined:
            meta.update(table_to_kv(t))

    subject_code = meta.get("Subject Code", "UNKNOWN")
    subject_name = meta.get("Syllabus Name") or meta.get("Course Name English", "Unknown")

    chunks = []

    # Chunk 0: tổng quan môn học (mô tả, điều kiện qua môn, tài liệu)
    overview_fields = ["Description", "NoCredit", "Pre-Requisite", "Time Allocation",
                        "StudentTasks", "Note", "MinAvgMarkToPass"]
    overview_text = "\n".join(f"{k}: {meta[k]}" for k in overview_fields if meta.get(k))
    if overview_text:
        chunks.append({
            "chunk_id": f"{subject_code}-overview",
            "subject_code": subject_code,
            "subject_name": subject_name,
            "section": "Overview & Grading Policy",
            "text": overview_text,
            "source_label": f"Syllabus {subject_code} — Overview & Grading Policy",
        })

    # Chunk mỗi CLO
    for c in clos:
        name = c.get("CLO Name", "").strip()
        detail = c.get("CLO Details", "").strip()
        if name and detail:
            chunks.append({
                "chunk_id": f"{subject_code}-{name}",
                "subject_code": subject_code,
                "subject_name": subject_name,
                "section": f"Learning Outcome {name}",
                "text": f"{name}: {detail}",
                "source_label": f"Syllabus {subject_code} — {name}",
            })

    # Chunk mỗi session (đây là nguồn chính cho deadline/task/tuần học)
    for s in sessions:
        session_no = s.get("Session", "").strip()
        topic = s.get("Topic", "").strip()
        tasks = s.get("Student's Tasks", "") or s.get("Student Tasks", "")
        materials = s.get("Student Materials", "")
        if not (session_no and topic):
            continue
        text = f"Session {session_no} — {topic}"
        if materials:
            text += f"\nTài liệu: {materials}"
        if tasks:
            text += f"\nNhiệm vụ sinh viên: {tasks}"
        chunks.append({
            "chunk_id": f"{subject_code}-session-{session_no}",
            "subject_code": subject_code,
            "subject_name": subject_name,
            "section": f"Session {session_no}",
            "text": text,
            "source_label": f"Syllabus {subject_code} — Session {session_no}",
        })

    return {
        "subject_code": subject_code,
        "subject_name": subject_name,
        "meta": meta,
        "clo_count": len(clos),
        "session_count": len(sessions),
        "chunks": chunks,
    }


if __name__ == "__main__":
    mode, path = sys.argv[1], sys.argv[2]
    out_path = None
    if "-o" in sys.argv:
        out_path = sys.argv[sys.argv.index("-o") + 1]

    if mode == "curriculum":
        result = parse_curriculum(path)
    elif mode == "syllabus":
        result = parse_syllabus(path)
    else:
        raise SystemExit("mode phải là 'curriculum' hoặc 'syllabus'")

    text = json.dumps(result, ensure_ascii=False, indent=2)
    if out_path:
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"Đã ghi {out_path}")
    else:
        print(text)
